from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from app.core.models import EmployeeProfile, Level, TrajectoryStep


def export_markdown(profile: EmployeeProfile, steps: list[TrajectoryStep], path: Path) -> None:
                                                           
    lines = _build_text_report(profile, steps)
    path.write_text("\n".join(lines), encoding="utf-8")


def export_csv(steps: list[TrajectoryStep], path: Path) -> None:
                                                      
    with path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.writer(file, delimiter=";")
        writer.writerow(["Код", "Компетенция", "Направление", "Текущий", "Целевой", "Приоритет", "Причина"])
        for step in steps:
            writer.writerow([
                step.competence.code,
                step.competence.title,
                step.competence.direction,
                step.current_level,
                step.target_level,
                f"{step.priority:.1f}",
                step.reason,
            ])


def export_json(profile: EmployeeProfile, steps: list[TrajectoryStep], path: Path) -> None:
                                                                 
    data = {
        "profile": profile.__dict__,
        "created_at": datetime.now().isoformat(timespec="minutes"),
        "steps": [
            {
                "code": step.competence.code,
                "title": step.competence.title,
                "direction": step.competence.direction,
                "current_level": step.current_level,
                "target_level": step.target_level,
                "priority": step.priority,
                "reason": step.reason,
            }
            for step in steps
        ],
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def export_docx(profile: EmployeeProfile, steps: list[TrajectoryStep], path: Path) -> None:
\
\
\
\
       
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as error:
        raise RuntimeError("Для выгрузки Word установите зависимость: pip install python-docx") from error

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = document.add_heading("Персональная траектория повышения цифровых компетенций", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"Сотрудник: {profile.full_name}")
    document.add_paragraph(f"Должность: {profile.position}")
    document.add_paragraph(f"Подразделение: {profile.department}")
    document.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    document.add_heading("Итоговая траектория", level=1)
    if not steps:
        document.add_paragraph("Все выбранные компетенции соответствуют целевому уровню.")
        document.save(path)
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["№", "Компетенция", "Направление", "Текущий", "Целевой", "Приоритет"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for number, step in enumerate(steps, start=1):
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = f"{step.competence.code} — {step.competence.title}"
        cells[2].text = step.competence.direction
        cells[3].text = Level.from_score(step.current_level).value
        cells[4].text = Level.from_score(step.target_level).value
        cells[5].text = f"{step.priority:.1f}/10"

    document.add_heading("Подробные рекомендации", level=1)
    for number, step in enumerate(steps, start=1):
        document.add_heading(f"{number}. {step.competence.code} — {step.competence.title}", level=2)
        document.add_paragraph(f"Описание: {step.competence.description}")
        document.add_paragraph(f"Причина рекомендации: {step.reason}")
        document.add_paragraph(f"Приоритет: {step.priority:.1f}/10")
        resources = document.add_paragraph("Учебные ресурсы:")
        resources.runs[0].bold = True
        for resource in step.competence.resources:
            document.add_paragraph(
                f"{resource.title} — {resource.kind}, {resource.duration_hours} ч.",
                style="List Bullet",
            )

    document.save(path)


def _build_text_report(profile: EmployeeProfile, steps: list[TrajectoryStep]) -> list[str]:
                                                                          
    lines = [
        "# Персональная траектория повышения цифровых компетенций",
        "",
        f"**Сотрудник:** {profile.full_name}",
        f"**Должность:** {profile.position}",
        f"**Подразделение:** {profile.department}",
        f"**Дата формирования:** {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
        "## Рекомендуемые шаги",
        "",
    ]
    for number, step in enumerate(steps, start=1):
        lines.extend(
            [
                f"### {number}. {step.competence.title}",
                f"- Направление: {step.competence.direction}",
                f"- Текущий уровень: {Level.from_score(step.current_level).value}",
                f"- Целевой уровень: {Level.from_score(step.target_level).value}",
                f"- Приоритет: {step.priority:.1f}/10",
                f"- Причина: {step.reason}",
                f"- Описание: {step.competence.description}",
                "- Ресурсы:",
            ]
        )
        for resource in step.competence.resources:
            lines.append(f"  - {resource.title} ({resource.kind}, {resource.duration_hours} ч.)")
        lines.append("")
    return lines
